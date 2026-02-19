from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os
import stat
import pandas as pd

class TankInvoiceGenerator:
    def delete_tables_in_first_page_header(self):
        """Delete all tables in the header of the first page only."""
        section = self.doc.sections[0]
        section.different_first_page_header = True
        first_page_header = section.first_page_header
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        first_header_element = first_page_header._element
        tables = first_header_element.findall('.//w:tbl', ns)
        for tbl in tables:
            tbl.getparent().remove(tbl)
            print(f"Deleted {len(tables)} table(s) from first page header.")
    def __init__(self, template_path="Template.docx"):
        """Initialize the document from template"""
        # Load the template document
        if os.path.exists(template_path):
            self.doc = Document(template_path)
            print(f"✓ Loaded template: {template_path}")
            
            # Remove any read-only protection inherited from template
            try:
                settings_element = self.doc.settings.element
                write_protection = settings_element.find(qn('w:writeProtection'))
                if write_protection is not None:
                    settings_element.remove(write_protection)
                doc_protection = settings_element.find(qn('w:documentProtection'))
                if doc_protection is not None:
                    settings_element.remove(doc_protection)
            except:
                pass  # Silently continue if protection removal fails during init
        else:
            print(f"⚠ Template not found at {template_path}, creating new document")
            self.doc = Document()
            
            # Set page margins to accommodate wider table
            sections = self.doc.sections
            for section in sections:
                section.left_margin = Inches(0.5)  # Reduced from default 1 inch
                section.right_margin = Inches(0.5)  # Reduced from default 1 inch
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
        
        self.table = None
        self.tanks = []
        self.gallon_type = "USG"  # Default
        self.template_path = template_path
        self.company_full_name = None  # Company name from company_details.xlsx (full_name)
        self.company_short_name = None  # Company brand name from company_details.xlsx (company_name)
        self.additional_details = []  # List of tuples: [(key, value), ...]
        self.note_section_gap = 6  # Default gap before NOTE section in points (adjustable)
        self.closing_paragraph_gap = 6  # Default gap before closing paragraph in points (adjustable)
        # Initialize total display flags (can be overridden by API)
        self.show_sub_total = True
        self.show_vat = True
        self.show_grand_total = True
        self.has_discount = False
        
    def get_user_inputs(self):
        """Collect all user inputs for tanks"""
        print("\n" + "="*60)
        print("WATER TANK INVOICE GENERATOR")
        print("="*60)
        
        # Get quotation header details
        print("\n" + "-"*60)
        print("QUOTATION HEADER DETAILS")
        print("-"*60)
        
        self.recipient_name = input("Recipient name (e.g., Ms. Hridya.): ").strip()
        self.recipient_company = input("Company name (e.g., M/s. Multiflags Contracting): ").strip()
        self.recipient_location = input("Location (e.g., Ajman, UAE): ").strip()
        
        # Phone with prefix
        phone_input = input("Phone number (e.g., + 971 50 312 8233): ").strip()
        self.recipient_phone = f"Phone: {phone_input}" if phone_input else ""
        
        # Email with prefix (optional)
        email_input = input("Email (optional, press Enter to skip): ").strip()
        self.recipient_email = f"Email: {email_input}" if email_input else ""
        
        # Auto-generate date in DD/MM/YY format from local computer
        self.quote_date = datetime.datetime.now().strftime("%d/%m/%y")
        print(f"Date (auto-generated): {self.quote_date}")
        
        self.quote_number = input("Quote number (e.g., GRPPT/2502/VV/2582): ").strip()
        
        self.subject = input("Subject (e.g., Supply and Installation of GRP Panel Water Tank): ").strip()
        self.project = input("Project location (e.g., Ajman.): ").strip()
        
        # Get number of tanks
        while True:
            try:
                num_tanks = int(input("\nNumber of types of tanks (sl.no.): "))
                if num_tanks > 0:
                    break
                print("Please enter a positive number.")
            except ValueError:
                print("Please enter a valid number.")
        
        # Get gallon type preference
        print("\nSelect gallon type for capacity display:")
        print("1. USG (US Gallons)")
        print("2. IMG (Imperial Gallons)")
        gallon_choice = input("Enter 1 or 2: ").strip()
        self.gallon_type = "IMG" if gallon_choice == "2" else "USG"
        
        # Collect data for each tank
        for n in range(1, num_tanks + 1):
            print("\n" + "-"*60)
            print(f"Tank number: {n}")
            print("-"*60)
            
            tank_data = {"sl_no": n}
            
            # Partition
            partition = input(f"With or without Partition of tank {n} (y/n): ").strip().lower()
            tank_data["partition"] = partition == "y"
            
            # Name
            tank_data["name"] = input(f"Name of tank {n}: ").strip()
            
            # Type
            print(f"\nType of tank {n}:")
            print("1. HOT PRESSED – NON - INSULATED")
            print("2. HOT PRESSED – INSULATED (BOTTOM & MANHOLE NON – INSULATED) – 5 side")
            print("3. HOT PRESSED – 6 SIDE INSULATED (BOTTOM & MANHOLE NON – INSULATED)")
            print("4. HOT PRESSED – 6 SIDE INSULATED (BOTTOM NON – INSULATED & MANHOLE)")
            print("5. Other (enter custom type)")
            
            type_choice = input("Enter choice (1-4): ").strip()
            types = {
                "1": "HOT PRESSED – NON - INSULATED",
                "2": "HOT PRESSED – INSULATED (BOTTOM & MANHOLE NON – INSULATED) – 5 side",
                "3": "HOT PRESSED – 6 SIDE INSULATED (BOTTOM & MANHOLE NON – INSULATED)",
                "4": "HOT PRESSED – 6 SIDE INSULATED (BOTTOM NON – INSULATED & MANHOLE)"
            }
            
            if type_choice in types:
                tank_data["type"] = types[type_choice]
            else:
                tank_data["type"] = input("Enter custom tank type: ").strip()
            
            # Size
            print(f"\nSize of tank {n}:")
            if tank_data["partition"]:
                print("(For partitioned tank, enter dimensions with partition notation, e.g., 2(1+1))")
                length_input = input("  L (in M): ").strip()
                tank_data["length_display"] = length_input
                # Extract actual length for calculation
                tank_data["length"] = self._parse_dimension(length_input)
                
                width_input = input("  W (in M): ").strip()
                tank_data["width_display"] = width_input
                # Extract actual width for calculation
                tank_data["width"] = self._parse_dimension(width_input)
            else:
                tank_data["length"] = float(input("  L (in M): "))
                tank_data["length_display"] = str(tank_data["length"])
                
                tank_data["width"] = float(input("  W (in M): "))
                tank_data["width_display"] = str(tank_data["width"])
            
            tank_data["height"] = float(input("  H (in M): "))
            
            # Calculate total capacity
            volume_m3 = tank_data["length"] * tank_data["width"] * tank_data["height"]
            tank_data["volume_m3"] = volume_m3
            
            if self.gallon_type == "USG":
                gallons = volume_m3 * 264.172  # Convert m3 to US gallons
            else:
                gallons = volume_m3 * 219.969  # Convert m3 to Imperial gallons
            
            tank_data["gallons"] = gallons
            
            # Free board (fixed at 30 cm = 0.3 M)
            tank_data["free_board"] = 0.3
            
            # Net volume calculation
            net_height = tank_data["height"] - tank_data["free_board"]
            tank_data["net_height"] = net_height
            
            # Determine skid type based on height
            height = tank_data["height"]
            if 2.0 <= height <= 3.0:
                tank_data["skid"] = "SKID BASE - HDG HOLLOW SECTION 50 X 50 X 3 MM (SQUARE TUBE)"
            elif 1.0 <= height < 1.5:
                tank_data["skid"] = "WITHOUT SKID"
            elif height > 3.0:
                tank_data["skid"] = "SKID BASE - I BEAM SKID"
            else:
                tank_data["skid"] = ""
            
            # Unit
            tank_data["unit"] = input(f"\nUnit of tank {n}: ").strip()
            
            # Quantity
            while True:
                try:
                    tank_data["qty"] = float(input(f"QTY of tank {n}: "))
                    break
                except ValueError:
                    print("Please enter a valid number.")
            
            # Unit price
            while True:
                try:
                    tank_data["unit_price"] = float(input(f"Unit price of tank {n} (AED): "))
                    break
                except ValueError:
                    print("Please enter a valid number.")
            
            # Total price
            tank_data["total_price"] = tank_data["qty"] * tank_data["unit_price"]
            
            self.tanks.append(tank_data)
        
        # Calculate total pages based on number of tanks (estimate: ~3-4 tanks per page)
        # This is a simple estimation; actual pages depend on content flow
        tanks_per_page = 3
        self.total_pages = max(1, (len(self.tanks) + tanks_per_page - 1) // tanks_per_page)
        self.quote_page = f"1/{self.total_pages}"  # For display purposes
        print(f"\nEstimated total pages: {self.total_pages}")
        print(f"Page numbering format: PAGE/TOTAL (will be dynamic)")
        
        # Check if ladder is needed
        self.needs_ladder = any(tank["height"] > 2.0 for tank in self.tanks)
        
        print("\n" + "="*60)
        print("All inputs collected successfully!")
        print("="*60)
        
        # Get gap size before NOTE section
        print("\n" + "-"*60)
        print("FORMATTING OPTIONS")
        print("-"*60)
        gap_input = input("Gap before NOTE section (in points, default 6, 0 for no gap): ").strip()
        if gap_input:
            try:
                self.note_section_gap = float(gap_input)
            except ValueError:
                print("Invalid input, using default (6 points)")
                self.note_section_gap = 6
        
        closing_gap_input = input("Gap before closing paragraph (in points, default 6, 0 for no gap): ").strip()
        if closing_gap_input:
            try:
                self.closing_paragraph_gap = float(closing_gap_input)
            except ValueError:
                print("Invalid input, using default (6 points)")
                self.closing_paragraph_gap = 6
        
        # Get additional sections configuration
        self._get_additional_sections_config()
    
    def _get_additional_sections_config(self):
        """Configure which additional sections to include"""
        print("\n" + "="*60)
        print("ADDITIONAL SECTIONS CONFIGURATION")
        print("="*60)
        
        self.sections = {
            'note': False,
            'closing': False,
            'signature': False,
            'material_spec': False,
            'warranty': False,
            'terms': False,
            'extra_note': True,
            'supplier_scope': False,
            'customer_scope': False,
            'final_note': False,
            'thank_you': True
        }
        
        # Ask for each section
        sections_prompt = [
            ('note', 'Include NOTE section (maintenance instructions)?'),
            ('closing', 'Include closing paragraph (We hope the above offer...)?'),
            ('signature', 'Include signature section (Yours truly, names)?'),
            ('material_spec', 'Include MATERIAL SPECIFICATION section?'),
            ('warranty', 'Include WARRANTY section?'),
            ('terms', 'Include TERMS AND CONDITIONS section?'),
            ('extra_note', 'Include extra NOTE section (deviations, delays, documents)?'),
            ('supplier_scope', 'Include SUPPLIER SCOPE section?'),
            ('customer_scope', 'Include CUSTOMER SCOPE section?'),
            ('final_note', 'Include final NOTE section?'),
            ('thank_you', 'Include THANK YOU FOR YOUR BUSINESS?')
        ]
        
        for key, prompt in sections_prompt:
            response = input(f"\n{prompt} (y/n): ").strip().lower()
            self.sections[key] = response == 'y'
        
        # Get content for enabled sections
        self._get_section_content()
    
    def _get_section_content(self):
        """Get and confirm content for each enabled section"""
        self.section_content = {}
        
        if self.sections['note']:
            self._configure_note_section()
        
        if self.sections['closing']:
            self._configure_closing_section()
        
        if self.sections['signature']:
            self._configure_signature_section()
        
        if self.sections['material_spec']:
            self._configure_material_spec_section()
        
        if self.sections['warranty']:
            self._configure_warranty_section()
        
        if self.sections['terms']:
            self._configure_terms_section()
        
        if self.sections['extra_note']:
            self._configure_extra_note_section()
        
        if self.sections['supplier_scope']:
            self._configure_supplier_scope_section()
        
        if self.sections['customer_scope']:
            self._configure_customer_scope_section()
        
        if self.sections['final_note']:
            self._configure_final_note_section()
    
    def _configure_note_section(self):
        """Configure NOTE section content"""
        print("\n" + "-"*60)
        print("NOTE SECTION")
        print("-"*60)
        
        default_notes = [
            "NOTE    :  1. DURING MAINTENANCE OF THE PARTITION TANK, THE WATER LEVELS IN EACH COMPARTMENT SHOULD BE REDUCED EQUALLY.",
            "                      THE MAXIMUM ALLOWABLE WATER HEIGHT IN EACH COMPARTMENT IS UPTO 1 MTR HEIGHT",
            "                  2. THE ABOVE TANK ONLY SUITABLE FOR STORING PORTABLE/ DRINKING WATER EXCEPT CHEMICAL /SOLID/TSE WATER.",
            "                  3. THE OFFER IS VALID FOR 30 DAYS FROM THE QUOTATION DATE."
        ]
        
        print("\nDefault NOTE content:")
        for note in default_notes:
            print(f"  {note}")
        
        response = input("\nUse default notes? (y/e/a): ").strip().lower()
        
        if response == 'y':
            self.section_content['note'] = default_notes
        elif response == 'e':
            notes = []
            for i, note in enumerate(default_notes, 1):
                print(f"\nNote {i}: {note}")
                edited = input("Edit this note (press Enter to keep, or type new text): ").strip()
                notes.append(edited if edited else note)
            self.section_content['note'] = notes
        else:  # add
            notes = list(default_notes)
            while True:
                new_note = input("\nEnter additional note (or press Enter to finish): ").strip()
                if not new_note:
                    break
                notes.append(new_note)
            self.section_content['note'] = notes
    
    def _configure_closing_section(self):
        """Configure closing paragraph"""
        default_text = ("We hope the above offer meets your requirements and awaiting the valuable order confirmation.\n"
                       "If you have any questions concerning the offer, please contact the undersigned.")
        
        print(f"\nDefault closing text:\n{default_text}")
        response = input("\nUse default? (y/e): ").strip().lower()
        
        if response == 'e':
            self.section_content['closing'] = input("Enter closing text: ").strip()
        else:
            self.section_content['closing'] = default_text
    
    def _configure_signature_section(self):
        """Configure signature section"""
        print("\n" + "-"*60)
        print("SIGNATURE SECTION")
        print("-"*60)
        
        # Ask for sales or office
        sig_type = input("\nSignatory type - Sales (s) or Office (o): ").strip().lower()
        
        if sig_type == 's':
            # Sales person - read from sales_person_details.xlsx
            try:
                sales_df = pd.read_excel('sales_person_details.xlsx')
                sales_names = sales_df['NAME'].tolist()
                
                print("\nAvailable Sales Persons:")
                for idx, name in enumerate(sales_names, 1):
                    print(f"{idx}. {name}")
                
                choice = int(input("\nSelect sales person number: ")) - 1
                selected_row = sales_df.iloc[choice]
                
                # Determine email column based on template
                if self.template_path.lower().endswith("grp_template.docx"):
                    email_col = 'EMAIL-GRPTANKS'
                elif self.template_path.lower().endswith("pipeco_template.docx"):
                    email_col = 'EMAIL-PIPECO'
                elif self.template_path.lower().endswith("colex_template.docx"):
                    email_col = 'EMAIL-COLEX'
                else:
                    email_col = 'EMAIL-GRPTANKS'  # Default
                
                # Get sales person details (left signatory)
                left_name = str(selected_row['NAME'])
                left_title = str(selected_row['DESIGNATION']) if 'DESIGNATION' in sales_df.columns else "Sales Executive"
                left_mobile = str(selected_row['MOB']) if 'MOB' in sales_df.columns else ""
                left_email = str(selected_row[email_col]) if email_col in sales_df.columns else ""
                
            except Exception as e:
                print(f"\n⚠ Error reading sales_person_details.xlsx: {e}")
                print("Using manual input instead.")
                left_name = input("Left signatory name (e.g., Viwin Varghese): ").strip() or "Viwin Varghese"
                left_title = input("Left signatory title (e.g., Sales Executive): ").strip() or "Sales Executive"
                left_mobile = input("Left signatory mobile (e.g., +971 54 450 4282): ").strip() or "+971 54 450 4282"
                left_email = input("Left signatory email (e.g., viwin@pipecogrp.com): ").strip() or "viwin@pipecogrp.com"
            
            # Right signatory (project manager) - read from Project_manager_details.xlsx
            try:
                pm_df = pd.read_excel('Project_manager_details.xlsx')
                
                if len(pm_df) == 1:
                    # Only one project manager, use default
                    selected_pm = pm_df.iloc[0]
                    print(f"\n✓ Using default Project Manager: {selected_pm['NAME']}")
                else:
                    pm_names = pm_df['NAME'].tolist()
                    print("\nAvailable Project Managers:")
                    for idx, name in enumerate(pm_names, 1):
                        print(f"{idx}. {name}")
                    
                    pm_choice = int(input("\nSelect project manager number: ")) - 1
                    selected_pm = pm_df.iloc[pm_choice]
                
                # Determine email column for project manager
                if self.template_path.lower().endswith("grp_template.docx"):
                    pm_email_col = 'EMAIL-GRPTANKS'
                elif self.template_path.lower().endswith("pipeco_template.docx"):
                    pm_email_col = 'EMAIL-PIPECO'
                elif self.template_path.lower().endswith("colex_template.docx"):
                    pm_email_col = 'EMAIL-COLEX'
                else:
                    pm_email_col = 'EMAIL-GRPTANKS'  # Default
                
                right_name = str(selected_pm['NAME'])
                right_title = str(selected_pm['DESIGNATION']) if 'DESIGNATION' in pm_df.columns else "Manager - Projects"
                right_mobile = str(selected_pm['MOB']) if 'MOB' in pm_df.columns else ""
                right_email = str(selected_pm[pm_email_col]) if pm_email_col in pm_df.columns else ""
                
            except Exception as e:
                print(f"\n⚠ Error reading Project_manager_details.xlsx: {e}")
                print("Using manual input instead.")
                right_name = input("Right signatory name (e.g., Anoop Mohan): ").strip() or "Anoop Mohan"
                right_title = input("Right signatory title (e.g., Manager - Projects): ").strip() or "Manager - Projects"
                right_mobile = input("Right signatory mobile (e.g., +971 50 952 4282): ").strip() or "+971 50 952 4282"
                right_email = input("Right signatory email (e.g., anoop@pipecogrp.com): ").strip() or "anoop@pipecogrp.com"
        
        else:  # office
            # Office - left signatory from Project_manager_details.xlsx
            try:
                pm_df = pd.read_excel('Project_manager_details.xlsx')
                
                if len(pm_df) == 1:
                    # Only one project manager, use default
                    selected_pm = pm_df.iloc[0]
                    print(f"\n✓ Using default Project Manager: {selected_pm['NAME']}")
                else:
                    pm_names = pm_df['NAME'].tolist()
                    print("\nAvailable Project Managers:")
                    for idx, name in enumerate(pm_names, 1):
                        print(f"{idx}. {name}")
                    
                    pm_choice = int(input("\nSelect project manager number: ")) - 1
                    selected_pm = pm_df.iloc[pm_choice]
                
                # Determine email column
                if self.template_path.lower().endswith("grp_template.docx"):
                    pm_email_col = 'EMAIL-GRPTANKS'
                elif self.template_path.lower().endswith("pipeco_template.docx"):
                    pm_email_col = 'EMAIL-PIPECO'
                elif self.template_path.lower().endswith("colex_template.docx"):
                    pm_email_col = 'EMAIL-COLEX'
                else:
                    pm_email_col = 'EMAIL-GRPTANKS'  # Default
                
                left_name = str(selected_pm['NAME'])
                left_title = str(selected_pm['DESIGNATION']) if 'DESIGNATION' in pm_df.columns else "Manager - Projects"
                left_mobile = str(selected_pm['MOB']) if 'MOB' in pm_df.columns else ""
                left_email = str(selected_pm[pm_email_col]) if pm_email_col in pm_df.columns else ""
                
            except Exception as e:
                print(f"\n⚠ Error reading Project_manager_details.xlsx: {e}")
                print("Using manual input instead.")
                left_name = input("Left signatory name (e.g., Anoop Mohan): ").strip() or "Anoop Mohan"
                left_title = input("Left signatory title (e.g., Manager - Projects): ").strip() or "Manager - Projects"
                left_mobile = input("Left signatory mobile (e.g., +971 50 952 4282): ").strip() or "+971 50 952 4282"
                left_email = input("Left signatory email (e.g., anoop@pipecogrp.com): ").strip() or "anoop@pipecogrp.com"
            
            # No right signatory for office
            right_name = ""
            right_title = ""
            right_mobile = ""
            right_email = ""
        
        # Auto-fetch signature image from signs&seals folder based on CODE
        signature_image = ""
        
        # Get CODE for signature images
        if left_name:
            try:
                # For sales, read from sales_person_details.xlsx; for office, read from Project_manager_details.xlsx
                if sig_type == 's':
                    df = pd.read_excel('sales_person_details.xlsx')
                else:
                    df = pd.read_excel('Project_manager_details.xlsx')
                
                code_row = df[df['NAME'] == left_name]
                if not code_row.empty:
                    code = code_row.iloc[0]['CODE']
                    # Look for signature image
                    for ext in ['.png', '.jpg', '.jpeg']:
                        sign_path = f"signs&seals/{code}_sign{ext}"
                        if os.path.exists(sign_path):
                            signature_image = sign_path
                            print(f"✓ Found signature image: {sign_path}")
                            break
                    if not signature_image:
                        print(f"⚠ Signature image not found for CODE: {code}")
            except Exception as e:
                print(f"⚠ Error fetching signature images: {e}")
        
        # Store signature content
        self.section_content['signature'] = {
            'left_name': left_name,
            'left_title': left_title,
            'left_mobile': left_mobile,
            'left_email': left_email,
            'right_name': right_name,
            'right_title': right_title,
            'right_mobile': right_mobile,
            'right_email': right_email,
            'signature_image': signature_image
        }
    
    def _configure_material_spec_section(self):
        """Configure material specification section"""
        default_specs = [
            "WRAS Approved Product.",
            "Sealant Tape – Non-Toxic PVC Foam Type.",
            "Roof Panel Vertical Support (Internal) – PVC Pipe.",
            "All Internal Metallic parts in continuous contact with water are Stainless Steel 316/A4 grade and External HDG Support Accessories with HDG Bolt/Nut/Washers.",
            "Manhole: 750mm Dia. with sealed cover and Lock. – 1 No.",
            "For 1 Mtr. height tank, Wall flat – 1 No. & Drain flat – 1 No.",
            "Clear Tube type level indicator (Without Marking) for 2 Mtr. height tank and above only.",
            "HDG Steel Skid with HDG Bolt / Nut / Washer for 2 Mtr. height tank and above only.",
            "Internal Ladder (HDG) and External Ladder (HDG) for 2 Mtr. height tank and above only.",
            "Air Vent, Inlet, Outlet, Overflow and Drain – 1 No. each with PVC flange (FL/PL) connections up to 3″.",
            "Manufacturer Warranty: 10 Year from the date of installation / testing and commissioning."
        ]
        
        print("\nDefault MATERIAL SPECIFICATION:")
        for spec in default_specs:
            print(f"  ➢   {spec}")
        
        response = input("\nUse default specs? (y/e): ").strip().lower()
        
        if response == 'e':
            specs = []
            for spec in default_specs:
                print(f"\n{spec}")
                edited = input("Edit (press Enter to keep): ").strip()
                specs.append(edited if edited else spec)
            self.section_content['material_spec'] = specs
        else:
            self.section_content['material_spec'] = default_specs
    
    def _configure_warranty_section(self):
        """Configure warranty section"""
        default_warranty = [
            "Any damage / loss caused directly or indirectly by natural calamities or any other force majeure conditions beyond the control of the supplier.",
            "Any damage occurs due to storing any chemicals, solids, or any other substances. (The proposed tank is specifically designed and intended only for potable / drinking water storage).",
            "Any defects or damage occur in the foundation that affect the tank.",
            "Any unauthorized modification or repairs made on the tank by parties other than the manufacturer representatives."
        ]
        
        print("\nDefault WARRANTY content:")
        for item in default_warranty:
            print(f"  ➢  {item}")
        
        response = input("\nUse default? (yes/edit): ").strip().lower()
        
        if response == 'edit':
            items = []
            for item in default_warranty:
                print(f"\n{item}")
                edited = input("Edit (press Enter to keep): ").strip()
                items.append(edited if edited else item)
            self.section_content['warranty'] = items
        else:
            self.section_content['warranty'] = default_warranty
    
    def _configure_terms_section(self):
        """Configure terms and conditions section"""
        default_terms = {
            'Price'   : 'The given prices are based on the supply and installation of the tank at your proposed site.',
            'Validity': 'The offer is valid for 30 days only.',
            'Delivery': 'One week from the receipt of advance payment.',
            'Payment' : 'Cash/CDC. 40% advance along with the confirmed order and 60% upon delivery of the material at the site. (In the event of late payment, a late payment charge of 2% per month on the contract value will be applied till the outstanding payment is settled).'
        }
        
        print("\nDefault TERMS AND CONDITIONS:")
        for key, value in default_terms.items():
            print(f"  ➢  {key:12} : {value}")
        
        response = input("\nUse default? (y/e): ").strip().lower()
        
        if response == 'e':
            terms = {}
            for key, value in default_terms.items():
                print(f"\n{key}: {value}")
                edited = input("Edit (press Enter to keep): ").strip()
                terms[key] = edited if edited else value
            self.section_content['terms'] = terms
        else:
            self.section_content['terms'] = default_terms
    
    def _configure_extra_note_section(self):
        """Configure extra NOTE section content"""
        print("\n" + "-"*60)
        print("EXTRA NOTE SECTION CONFIGURATION")
        print("-"*60)
        
        company_name = self._get_company_name()
        
        default_notes = [
            "Any deviations from this quotation to suit the site's condition will have additional cost implications.",
            "If the work is indefinitely delayed beyond 30 days after the delivery of materials due to the issues caused by the customer or site condition, the Company will not be liable for any damage to the supplied materials.",
            "The submission of all related documents, including the warranty certificate, will be done upon receiving the final payment.",
            "Any additional test / lab charges incurred from third parties / external agencies are under the scope of the contractor / client.",
            f"Until receiving the final settlement from the client, {company_name} has reserved the right to use the supplied materials at the site.",
            "The testing and commissioning should be completed within a period of 15 to 30 days from the installation completion date by the Contractor/Client.",
            "For the net volume, a minimum of 30 cm freeboard area is to be calculated from the total height of the tank."
        ]
        
        print("\nDefault extra NOTE items:")
        for idx, item in enumerate(default_notes, 1):
            print(f"{idx}. {item}")
        
        response = input("\nUse default items? (y/n): ").strip().lower()
        if response == 'y':
            self.section_content['extra_note'] = default_notes
        else:
            print("\nEnter custom extra NOTE items (one per line, empty line to finish):")
            items = []
            while True:
                item = input(f"Item {len(items) + 1}: ").strip()
                if not item:
                    break
                items.append(item)
            self.section_content['extra_note'] = items if items else default_notes
    
    def _configure_supplier_scope_section(self):
        """Configure supplier scope section"""
        default_scope = [
            "Supply, installation & supervision for T & C of the tank at the site.",
            "Basic Hand Toolbox.",
            "Surveyors Equipment's for base skid levelling.",
            "Power Tools – Welding/Grinder/Drill/Tighter Machine /Cables.",
            "Flanges as mentioned in the offer."
        ]
        
        print("\nDefault SUPPLIER SCOPE:")
        for item in default_scope:
            print(f"  ➢  {item}")
        
        response = input("\nUse default? (yes/edit): ").strip().lower()
        
        if response == 'edit':
            items = []
            for item in default_scope:
                print(f"\n{item}")
                edited = input("Edit (press Enter to keep): ").strip()
                items.append(edited if edited else item)
            self.section_content['supplier_scope'] = items
        else:
            self.section_content['supplier_scope'] = default_scope
    
    def _configure_customer_scope_section(self):
        """Configure customer scope section"""
        default_scope = [
            "Material offloading, safe storage, and shifting near the foundation. (If the offloading and lifting team is not ready upon our vehicle's arrival, the delivery may be rescheduled, and a maximum charge of AED 1000 will be applied to the customer for rescheduling).",
            "Crain/Boom Loader/other facilities for offloading.",
            "Other plumbing works / Float Valves/ Valves / Float Switches.",
            "Scaffolding as per the site condition.",
            "Flanges other than specified.",
            "Water Thermos & Rest Rooms to be provided by the Contractor/Client.",
            "Electricity/Generator for installation and water for testing.",
            "Accommodation for the technical staff should be provided by the client/contractor.",
            "Grouting, if required for levelling the tank foundation. (After completing the skid work for clearing the space between the base skid and plinth). After the grouting process, a minimum of 3 days will be required to schedule the installation.",
            "In case of any leakage detected after filling the tank, it shall be drained out (if required) and bear any related expenses for refilling the tank for retesting after the rectification.",
            "Any obligations, including entry permits, labour passes and risk liability insurance policy charges etc."
        ]
        
        print("\nDefault CUSTOMER SCOPE:")
        for item in default_scope:
            print(f"  ➢  {item}")
        
        response = input("\nUse default? (yes/edit): ").strip().lower()
        
        if response == 'edit':
            items = []
            for item in default_scope:
                print(f"\n{item}")
                edited = input("Edit (press Enter to keep): ").strip()
                items.append(edited if edited else item)
            self.section_content['customer_scope'] = items
        else:
            self.section_content['customer_scope'] = default_scope
    
    def _configure_final_note_section(self):
        """Configure final note section"""
        company_name = self._get_company_name()
        default_notes = [
            "Any deviations from this quotation to suit the site's condition will have additional cost implications.",
            "If the work is indefinitely delayed beyond 30 days after the delivery of materials due to the issues caused by the customer or site condition, the company will not be liable for any damage to the supplied materials.",
            "The submission of all related documents, including the warranty certificate, will be done upon receiving the final payment.",
            "Any additional test / lab charges incurred from third parties / external agencies are under the scope of the contractor / client.",
            f"Until receiving the final settlement from the client, {company_name} has reserved the right to use the supplied materials at the site.",
            "The testing and commissioning should be completed within a period of 15 to 30 days from the installation completion date by the Contractor/Client.",
            "For the net capacity, a minimum of 30 cm freeboard area is to be calculated from the total height of the tank."
        ]
        
        print("\nDefault final NOTE:")
        for note in default_notes:
            print(f"  ➢  {note}")
        
        response = input("\nUse default? (yes/edit): ").strip().lower()
        
        if response == 'edit':
            notes = []
            for note in default_notes:
                print(f"\n{note}")
                edited = input("Edit (press Enter to keep): ").strip()
                notes.append(edited if edited else note)
            self.section_content['final_note'] = notes
        else:
            self.section_content['final_note'] = default_notes
    
    def _parse_dimension(self, dim_str):
        """Parse dimension string like '2(1+1)' to get actual value"""
        try:
            # Try to evaluate simple expressions like 2(1+1) -> 2
            # Remove spaces
            dim_str = dim_str.replace(" ", "")
            # If it contains parentheses, extract the coefficient
            if "(" in dim_str:
                # Extract number before parentheses
                return float(dim_str.split("(")[0])
            else:
                return float(dim_str)
        except:
            return float(dim_str)
    
    def _get_company_name(self):
        """Get company name from frontend or based on template"""
        # If company name was set from frontend (via API), use it
        if self.company_full_name:
            return self.company_full_name
        
        # Fallback to template-based detection for backward compatibility
        if self.template_path.lower().endswith("pipeco_template.docx"):
            return "GRP PIPECO TANKS TRADING L.L.C"
        elif self.template_path.lower().endswith("colex_template.docx"):
            return "COLEX TANKS TRADING L.L.C"
        elif self.template_path.lower().endswith("grp_template.docx"):
            return "GRP TANKS TRADING L.L.C"
        else:
            return "GRP PIPECO TANKS TRADING L.L.C"  # Default
    
    def _find_common_elements(self):
        """Find common elements across all tanks"""
        common_elements = []
        
        # Check if all tanks have the same type
        if len(self.tanks) > 0:
            first_type = self.tanks[0]["type"]
            if all(tank["type"] == first_type for tank in self.tanks):
                common_elements.append(("type", first_type))
        
        # Check if all tanks have the same skid
        if len(self.tanks) > 0:
            first_skid = self.tanks[0]["skid"]
            if first_skid and all(tank["skid"] == first_skid for tank in self.tanks):
                common_elements.append(("skid", first_skid))
        
        return common_elements
    
    def _get_common_support_system(self):
        """Check if all tanks have the same support system"""
        if len(self.tanks) == 0:
            return None
        
        first_support = self.tanks[0].get("support_system", "Internal")
        if all(tank.get("support_system", "Internal") == first_support for tank in self.tanks):
            return first_support
        return None  # Mixed support systems
    
    def _get_support_system_text(self, support_type):
        """Get the description text for a support system type"""
        if support_type == "External":
            return "EXTERNAL REINFORCEMENT SYSTEM"
        else:  # Internal or default
            return "INTERNAL SS 316 AND EXTERNAL HDG SUPPORT SYSTEM"

    
    def _add_header_to_all_pages(self):
        """Add quote box ONLY to pages 2+ header, keep first page header with template elements only"""
        from copy import deepcopy
        
        section = self.doc.sections[0]
        
        # ═══════════════════════════════════════════════════════════════════════════════
        # CRITICAL: Use XML-level manipulation to ensure truly separate headers
        # ═══════════════════════════════════════════════════════════════════════════════
        
        # Step 1: Enable different first page header
        section.different_first_page_header = True
        
        # Step 2: Get the XML element for the section properties
        sectPr = section._sectPr
        
        # Step 3: Get header references
        # In OOXML, headers are referenced by headerReference elements in sectPr
        # We need to ensure first page header and default header are different
        
        # Get the default header (for pages 2+)
        default_header = section.header
        
        # Get the first page header
        first_page_header = section.first_page_header
        
        # ═══════════════════════════════════════════════════════════════════════════════
        # Step 4: CLEAR first page header of ALL tables (quote boxes) at XML level
        # ═══════════════════════════════════════════════════════════════════════════════
        first_header_element = first_page_header._element
        # Find and remove ALL w:tbl elements (tables) from first page header
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for tbl in first_header_element.findall('.//w:tbl', ns):
            tbl.getparent().remove(tbl)
        print(f"✓ FIRST PAGE header: Cleared all tables at XML level")
        
        # ═══════════════════════════════════════════════════════════════════════════════
        # Step 5: Add quote box to DEFAULT header (pages 2+) ONLY
        # ═══════════════════════════════════════════════════════════════════════════════
        self._add_quote_box_to_header(default_header)
        
        # ═══════════════════════════════════════════════════════════════════════════════
        # Step 6: FINAL CLEANUP - ensure first page header has NO tables
        # ═══════════════════════════════════════════════════════════════════════════════
        # Do this AGAIN after adding quote box to default header, in case of any linking
        first_header_element = first_page_header._element
        tables_found = first_header_element.findall('.//w:tbl', ns)
        for tbl in tables_found:
            tbl.getparent().remove(tbl)
        
        print(f"✓ FIRST PAGE header final: {len(first_page_header.tables)} table(s), {len(first_page_header.paragraphs)} paragraph(s)")
        print(f"✓ PAGES 2+ header: {len(default_header.tables)} table(s) (should be 1 for quote box)")
    
    # Seal image insertion removed - no longer adding seal to footers
    
    def _add_quote_box_to_header(self, header):
        """Add the quote box table to a specific header"""
        # Add spacer for some templates
        if len(header.paragraphs) > 0 or len(header.tables) > 0:
            if self.template_path.lower().endswith("colex_template.docx"):
                # Add multiple line breaks to push table below header content
                for _ in range(3):
                    spacer = header.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(8)
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.line_spacing = 1.0
                    spacer.paragraph_format.left_indent = Pt(0)
                    spacer.paragraph_format.right_indent = Pt(0)
                    spacer.paragraph_format.first_line_indent = Pt(0)
        
        
        # Box width based on template
        if self.template_path.lower().endswith("colex_template.docx"):
            box_width = Inches(8)
        elif self.template_path.lower().endswith("grp_template.docx"):
            box_width = Inches(8)
        elif self.template_path.lower().endswith("pipeco_template.docx"):
            box_width = Inches(8.5)
        else:
            box_width = Inches(8)
        
        # Create quote box table
        quote_table = header.add_table(rows=1, cols=3, width=box_width)
        quote_table.autofit = False
        
        # Position table based on template
        if self.template_path.lower().endswith("colex_template.docx"):
            tbl = quote_table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            tblInd = tblPr.find(qn('w:tblInd'))
            if tblInd is not None:
                tblPr.remove(tblInd)
            tblInd = OxmlElement('w:tblInd')
            tblInd.set(qn('w:w'), '-360')
            tblInd.set(qn('w:type'), 'dxa')
            tblPr.append(tblInd)
            
            jc = tblPr.find(qn('w:jc'))
            if jc is not None:
                tblPr.remove(jc)
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'left')
            tblPr.append(jc)
        elif self.template_path.lower().endswith("grp_template.docx") or self.template_path.lower().endswith("pipeco_template.docx"):
            tbl = quote_table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            tblCellSpacing = OxmlElement('w:tblCellSpacing')
            tblCellSpacing.set(qn('w:w'), '0')
            tblCellSpacing.set(qn('w:type'), 'dxa')
            tblPr.append(tblCellSpacing)
        
        # Calculate cell widths
        quote_text = f'QUOTE NO : {self.quote_number}'
        date_text = f'DATE : {self.quote_date}'
        page_text = f'PAGE NO : {self.quote_page}'
        
        chars_per_inch = 9
        quote_width = max(1.5, len(quote_text) / chars_per_inch)
        date_width = max(1.5, len(date_text) / chars_per_inch)
        page_width = max(1.5, len(page_text) / chars_per_inch)
        
        total_available = 7.5
        total_needed = quote_width + date_width + page_width
        
        if total_needed > total_available:
            scale = total_available / total_needed
            quote_width *= scale
            date_width *= scale
            page_width *= scale
        else:
            extra = total_available - total_needed
            quote_width += extra * (quote_width / total_needed)
            date_width += extra * (date_width / total_needed)
            page_width += extra * (page_width / total_needed)
        
        # Table borders
        tbl = quote_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Remove cell padding
        tcMar = OxmlElement('w:tblCellMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tblPr.append(tcMar)
        
        sky_blue = RGBColor(20, 123, 197)
        row = quote_table.rows[0]
        
        # Cell 1: QUOTE NO
        cell = row.cells[0]
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(quote_width * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        cell._element.get_or_add_tcPr().append(tcW)
        
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(f'QUOTE NO : ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        run = para.add_run(self.quote_number)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        
        # Cell 2: DATE
        cell = row.cells[1]
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(date_width * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        cell._element.get_or_add_tcPr().append(tcW)
        
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(f'DATE : ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        run = para.add_run(self.quote_date)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        
        # Cell 3: PAGE NO
        cell = row.cells[2]
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(page_width * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        cell._element.get_or_add_tcPr().append(tcW)
        
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(f'PAGE NO : ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        
        run = para.add_run()
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        self._add_page_number_field(run)

    def _add_page_number_field(self, run, use_blue_color=True):
        """Add a dynamic PAGE/NUMPAGES field to a run for automatic page numbering"""
        # Store existing font properties
        existing_bold = run.font.bold
        existing_color = run.font.color.rgb if run.font.color.rgb else None
        
        # Create the field code elements for PAGE (current page)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        # Add placeholder text for PAGE
        t1 = OxmlElement('w:t')
        t1.text = '1'
        
        fldChar2End = OxmlElement('w:fldChar')
        fldChar2End.set(qn('w:fldCharType'), 'end')
        
        # Add field elements to the run
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(t1)
        run._r.append(fldChar2End)
        
        # Add the "/" separator
        t = OxmlElement('w:t')
        t.text = '/'
        run._r.append(t)
        
        # Create field code elements for NUMPAGES (total pages)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = 'NUMPAGES'
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'separate')
        
        # Add placeholder text for NUMPAGES
        t2 = OxmlElement('w:t')
        total_pages_value = getattr(self, 'total_pages', 1)
        t2.text = str(total_pages_value)
        
        fldChar4End = OxmlElement('w:fldChar')
        fldChar4End.set(qn('w:fldCharType'), 'end')
        
        # Add field elements for total pages
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)
        run._r.append(t2)
        run._r.append(fldChar4End)
        
        # Set font properties for the page number
        if use_blue_color:
            run.font.bold = True
            run.font.color.rgb = RGBColor(20, 123, 197)  # Sky blue
        # If use_blue_color is False, keep existing formatting
    
    def _add_quote_box_to_body(self):
        """Add quote information box to document body (not header)"""
        # Create a 1-row, 3-column table for the quote box
        quote_table = self.doc.add_table(rows=1, cols=3)
        quote_table.autofit = False
        
        # Prepare content for cells to calculate optimal widths
        quote_text = f'QUOTE NO : {self.quote_number}'
        date_text = f'DATE : {self.quote_date}'
        page_text = 'PAGE NO : 1/3'  # First page always shows 1/3
        
        # Calculate approximate character widths for Calibri 11 Bold (roughly 9 chars per inch)
        chars_per_inch = 9
        quote_width = max(1.5, len(quote_text) / chars_per_inch)
        date_width = max(1.5, len(date_text) / chars_per_inch)
        page_width = max(1.5, len(page_text) / chars_per_inch)
        
        # Total available width (approximately 7.5 inches for standard page with margins)
        total_available = 7.5
        total_needed = quote_width + date_width + page_width
        
        # Scale proportionally if needed
        if total_needed > total_available:
            scale = total_available / total_needed
            quote_width *= scale
            date_width *= scale
            page_width *= scale
        else:
            # If text is shorter, distribute extra space proportionally
            extra_space = total_available - total_needed
            quote_width += extra_space * (quote_width / total_needed)
            date_width += extra_space * (date_width / total_needed)
            page_width += extra_space * (page_width / total_needed)
        
        # Apply borders to the table
        tbl = quote_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Add table borders
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Remove cell padding
        tcMar = OxmlElement('w:tblCellMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tblPr.append(tcMar)
        
        # Sky blue color: RGB(20, 123, 197)
        sky_blue = RGBColor(20, 123, 197)
        
        # Fill the cells
        row = quote_table.rows[0]
        
        # Cell 1: QUOTE NO
        cell = row.cells[0]
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(quote_width * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        cell._element.get_or_add_tcPr().append(tcW)
        
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(f'QUOTE NO : ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        run = para.add_run(self.quote_number)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        
        # Cell 2: DATE
        cell = row.cells[1]
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(date_width * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        cell._element.get_or_add_tcPr().append(tcW)
        
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(f'DATE : ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        run = para.add_run(self.quote_date)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        
        # Cell 3: PAGE NO
        cell = row.cells[2]
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(page_width * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        cell._element.get_or_add_tcPr().append(tcW)
        
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(f'PAGE NO : ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        
        # Add static page number "1/3" for first page
        run = para.add_run('1/3')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        
        # Add spacing after the quote box
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(6)
    
    def _add_header_table_for_additional_pages(self):
        """Add a header table with quote info for pages 2 onwards"""
        # Create a 1-row, 3-column table for the header
        header_table = self.doc.add_table(rows=1, cols=3)
        header_table.autofit = False
        
        # Prepare content for cells to calculate optimal widths
        quote_text = f'QUOTE NO : {self.quote_number}'
        date_text = f'DATE : {self.quote_date}'
        page_text = f'PAGE NO : {self.quote_page}'
        
        # Calculate approximate character widths for Calibri 11 Bold (roughly 9 chars per inch)
        # Using 9 characters per inch for bold text
        chars_per_inch = 9
        quote_width = max(1.5, len(quote_text) / chars_per_inch)
        date_width = max(1.5, len(date_text) / chars_per_inch)
        page_width = max(1.5, len(page_text) / chars_per_inch)
        
        # Total available width (approximately 7.5 inches for standard page with margins)
        total_available = 7.5
        total_needed = quote_width + date_width + page_width
        
        # Scale proportionally if needed
        if total_needed > total_available:
            scale_factor = total_available / total_needed
            quote_width *= scale_factor
            date_width *= scale_factor
            page_width *= scale_factor
        else:
            # Distribute remaining space proportionally
            remaining = total_available - total_needed
            quote_width += remaining * (quote_width / total_needed)
            date_width += remaining * (date_width / total_needed)
            page_width += remaining * (page_width / total_needed)
        
        # Apply borders to the table
        tbl = header_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Add table borders
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Remove cell padding
        tcMar = OxmlElement('w:tblCellMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '30')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tblPr.append(tcMar)
        
        # Sky blue color: RGB(135, 206, 235)
        sky_blue = RGBColor(20, 123, 197)
        
        # Fill the cells
        row = header_table.rows[0]
        
        # Cell 1: QUOTE NO
        cell = row.cells[0]
        # Set cell width using XML (convert inches to twips: 1 inch = 1440 twips)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(quote_width * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        cell._element.get_or_add_tcPr().append(tcW)
        
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(f'QUOTE NO : ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        run = para.add_run(self.quote_number)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        
        # Cell 2: DATE
        cell = row.cells[1]
        # Set cell width using XML
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(date_width * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        cell._element.get_or_add_tcPr().append(tcW)
        
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(f'DATE : ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        run = para.add_run(self.quote_date)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        
        # Cell 3: PAGE NO
        cell = row.cells[2]
        # Set cell width using XML
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(page_width * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        cell._element.get_or_add_tcPr().append(tcW)
        
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(f'PAGE NO : ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = sky_blue
        
        # Add dynamic page number field (current/total)
        run = para.add_run()
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        self._add_page_number_field(run)
        
        # Add spacing after the table
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(6)

    def create_invoice_table(self):
        """Create the invoice table with all tank details"""
        # Add header with quote box and green line - will appear on ALL pages
        self._add_header_to_all_pages()
        # Ensure all tables are deleted from first page header (per user request)
        self.delete_tables_in_first_page_header()
        
        # Seal insertion removed - no longer adding seal to footers
        
        # Add quotation header before table
        self._create_quotation_header()
        
        # Calculate number of rows needed
        # Header + common row + tank rows + footer rows (based on which totals are enabled)
        show_sub = getattr(self, 'show_sub_total', True)
        show_vat = getattr(self, 'show_vat', True)
        show_grand = getattr(self, 'show_grand_total', True)
        num_footer_rows = sum([show_sub, show_vat, show_grand])
        
        num_tank_rows = len(self.tanks)
        total_rows = 1 + 1 + num_tank_rows + num_footer_rows
        
        # Add a paragraph break before the table if template has content
        # Then immediately remove it to avoid Aptos font gap
        if len(self.doc.paragraphs) > 0:
            para = self.doc.add_paragraph()
            # Remove spacing to prevent table from being pushed down
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            # Remove the paragraph element to prevent Aptos font gap
            p = para._element
            p.getparent().remove(p)
        
        # Create table and add it to the document
        self.table = self.doc.add_table(rows=total_rows, cols=6)
        
        # Try to apply Table Grid style, fallback to manual borders if not available
        try:
            self.table.style = 'Table Grid'
        except KeyError:
            # Template doesn't have Table Grid style, manually add borders
            print("⚠ 'Table Grid' style not found in template, applying manual borders")
            self._apply_table_borders()
        
        # Remove cell padding to make text compact like image 1
        self._remove_cell_padding()
        
        # Set column widths - UNIT and QTY increased to fit "GRAND TOTAL" on single line
        widths = [Inches(0.6), Inches(9.625), Inches(1.1), Inches(1.0), Inches(1.4), Inches(1.4)]
        for idx, width in enumerate(widths):
            for cell in self.table.columns[idx].cells:
                cell.width = width
        
        # Format header row (row 0)
        self._create_header()
        
        # Row 1: Common information row (merge all columns)
        self._merge_cells(1, 0, 1, 5)
        self._fill_common_row()
        
        # Add tank rows (starting from row 2)
        row_idx = 2
        tank_start_rows = {}  # Track starting row for each sl_no
        
        for tank in self.tanks:
            sl_no = tank['sl_no']
            option_num = tank.get('option_number', 1)
            
            # Track the first row for this sl_no
            if sl_no not in tank_start_rows:
                tank_start_rows[sl_no] = row_idx
            
            self._fill_tank_row(row_idx, tank)
            row_idx += 1
        
        # Merge SL. NO. cells for tanks with multiple options
        for sl_no, start_row in tank_start_rows.items():
            # Find all tanks with this sl_no
            tanks_with_sl = [t for t in self.tanks if t['sl_no'] == sl_no]
            if len(tanks_with_sl) > 1:
                # Merge from start_row to start_row + len(tanks_with_sl) - 1
                end_row = start_row + len(tanks_with_sl) - 1
                self._merge_cells(start_row, 0, end_row, 0)
        
        # Footer rows
        footer_start = row_idx
        self._create_footer(footer_start)
        
        # Apply column-specific padding
        self._apply_column_specific_padding()
        
        # Apply Calibri 11 to all cells
        self._apply_font_to_all_cells()
        
        # Add additional sections after the table (these will flow across all pages)
        self._add_additional_sections()
    
    def _create_quotation_header(self):
        """Create quotation header content above the table"""
        # CRITICAL FIX: Remove ALL paragraphs from document body to eliminate any gaps
        # This ensures we start with a completely clean slate
        while len(self.doc.paragraphs) > 0:
            para = self.doc.paragraphs[0]  # Always remove first paragraph
            p = para._element
            p.getparent().remove(p)
        
        # ADJUSTED: Line ~281 - QUOTATION title with ABSOLUTELY NO spacing above
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Remove ALL spacing - above, below, and line spacing
        fmt = title.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.0
        fmt.line_spacing_rule = 1  # Single line spacing rule
        
        # Also remove any inherited spacing
        title._element.get_or_add_pPr()
        pPr = title._element.pPr
        # Remove spacing element if exists
        for spacing in pPr.findall(qn('w:spacing')):
            pPr.remove(spacing)
        # Add explicit zero spacing
        spacing_elem = OxmlElement('w:spacing')
        spacing_elem.set(qn('w:before'), '0')
        spacing_elem.set(qn('w:after'), '0')
        spacing_elem.set(qn('w:line'), '240')  # 240 twips = 12pt single spacing
        spacing_elem.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing_elem)
        
        run = title.add_run('QUOTATION')
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = RGBColor(0, 32, 96)  # Color #002060
        run.underline = True
        
        # "To." text outside table
        to_para = self.doc.add_paragraph()
        to_para.paragraph_format.space_before = Pt(0)
        to_para.paragraph_format.space_after = Pt(0)
        to_para.paragraph_format.line_spacing = 1.0
        run = to_para.add_run('To.')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Create table with 1 row and 2 columns for side-by-side layout
        info_table = self.doc.add_table(rows=1, cols=2)
        info_table.autofit = False
        # ⚠️ ADJUST HERE TO MOVE VERTICAL BORDER ⚠️
        # To move vertical border RIGHT: increase left value (e.g., 4.5, 5.0, 5.5)
        # Current: 5.0 (left) + 2.5 (right) = 7.5 inches total width
        info_table.columns[0].width = Inches(5.0)  # Left column - ✅ ADJUST THIS TO MOVE BORDER
        info_table.columns[1].width = Inches(2.5)  # Right column - ✅ ADJUST THIS
        
        
        # Add left indent to move table to the right
        tbl = info_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Set table indent
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '360')  # 360 twips = 0.25 inches indent
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        
        # Remove existing borders
        for borders in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(borders)
        
        # Add borders set to 'none'
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Remove cell borders and padding
        for row in info_table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                
                # Remove existing borders
                for tcBorders in tcPr.findall(qn('w:tcBorders')):
                    tcPr.remove(tcBorders)
                
                # Add cell borders set to 'none'
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
                
                # Remove cell padding
                for tcMar in tcPr.findall(qn('w:tcMar')):
                    tcPr.remove(tcMar)
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '0')
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)
        
        # LEFT CELL: Recipient information
        left_cell = info_table.rows[0].cells[0]
        # Remove default empty paragraph that has Aptos font
        if len(left_cell.paragraphs) > 0:
            p = left_cell.paragraphs[0]._element
            p.getparent().remove(p)
        
        # Recipient name (only add if not empty)
        if self.recipient_name:
            para = left_cell.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(self.recipient_name)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
        
        # Company name (only add if not empty)
        if self.recipient_company:
            para = left_cell.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(self.recipient_company)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
        
        # Location (only add if not empty)
        if self.recipient_location:
            para = left_cell.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(self.recipient_location)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
        
        # Phone (only add if not empty)
        if self.recipient_phone and self.recipient_phone.strip():
            para = left_cell.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            
            # Check if phone has colon and format accordingly
            phone_text = self.recipient_phone.strip()
            if ':' in phone_text:
                # Add with colon on same line
                run = para.add_run(phone_text)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.bold = True
            else:
                # No colon, add as-is
                run = para.add_run(phone_text)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.bold = True
        
        # Email (only add if not empty)
        if self.recipient_email and self.recipient_email.strip():
            para = left_cell.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            
            # Check if email has "Email:" prefix and format accordingly
            email_text = self.recipient_email.strip()
            if ':' in email_text:
                # Add with colon on same line
                run = para.add_run(email_text)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.bold = True
            else:
                # No colon, add as-is
                run = para.add_run(email_text)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.bold = True
        
        # Set left cell width using XML (7200 twips = 5.0 inches)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), '6900')  # ⚠️ ADJUST: 1440 twips = 1 inch, so 7200 = 5.0"
        tcW.set(qn('w:type'), 'dxa')
        left_cell._element.get_or_add_tcPr().append(tcW)
        
        # RIGHT CELL: Quote information
        right_cell = info_table.rows[0].cells[1]
        # Remove default empty paragraph that has Aptos font
        if len(right_cell.paragraphs) > 0:
            p = right_cell.paragraphs[0]._element
            p.getparent().remove(p)
        
        # Set vertical alignment to top
        tcPr = right_cell._element.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'top')
        tcPr.append(vAlign)
        
        # Set right cell width using XML (3600 twips = 2.5 inches)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), '3900')  # ⚠️ ADJUST: 1440 twips = 1 inch, so 3600 = 2.5"
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        
        # Date (always DD/MM/YYYY)
        para = right_cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        # Add tab stops for colon alignment like Subject/Project
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(1.2))  # Position for colon
        # Format date as DD/MM/YYYY
        date_str = self.quote_date
        # If date is DD/MM/YY, convert to DD/MM/YYYY
        if isinstance(date_str, str) and len(date_str.split('/')) == 3:
            d, m, y = date_str.split('/')
            if len(y) == 2:
                y = '20' + y
            date_str = f"{d}/{m}/{y}"
        # Add Date with value on same line
        run = para.add_run('Date            : {}'.format(date_str))
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Page - Keep blank (no page number displayed)
        para = right_cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        # Add tab stops for colon alignment
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(1.2))  # Position for colon
        
        # Add Page with value on same line
        run = para.add_run('Page            : ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Add the page number field to the same run
        self._add_page_number_field(run, use_blue_color=False)
        
        # Quote No
        para = right_cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        # Add tab stops for colon alignment
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(1.2))  # Position for colon
        # Add Quote No with value on same line
        run = para.add_run('Quote No.  : {}'.format(self.quote_number))
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Spacing after table
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)
        
        # LOCATION ADJUSTED: Line ~486 - Subject line with reduced space above
        # Subject line (vertically aligned with "To.")
        subject_para = self.doc.add_paragraph()
        subject_para.paragraph_format.space_before = Pt(0)  # No space above subject
        subject_para.paragraph_format.space_after = Pt(0)
        # Add Subject with value on same line
        run = subject_para.add_run('Subject          : ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run2 = subject_para.add_run(self.subject)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(10)
        run2.font.bold = True
        run2.underline = True
        
        # Project line
        project_para = self.doc.add_paragraph()
        project_para.paragraph_format.space_before = Pt(0)
        project_para.paragraph_format.space_after = Pt(0)  # Remove gap
        # Add Project with value on same line
        run = project_para.add_run('Project           : ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run2 = project_para.add_run(self.project)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(10)
        run2.font.bold = True
        run2.underline = True
        
        # Additional details - Display like Subject and Project (simple paragraphs)
        if hasattr(self, 'additional_details') and self.additional_details:
            # Filter out completely empty entries (where both key and value are empty)
            valid_details = [
                (key or "", value or "") 
                for key, value in self.additional_details 
                if key or value  # Include if at least key or value exists
            ]
            
            # Add each additional detail as a separate line matching Subject/Project format
            for key, value in valid_details:
                detail_para = self.doc.add_paragraph()
                detail_para.paragraph_format.space_before = Pt(0)
                detail_para.paragraph_format.space_after = Pt(0)
                
                # Add key with padding to align with Subject/Project (18 chars total width)
                key_text = key.strip() if key else ""
                padding_needed = max(0, 18 - len(key_text))
                padded_key = key_text + ' ' * padding_needed + ': '
                
                run = detail_para.add_run(padded_key)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.bold = True
                
                # Add value (bold and underlined like Subject/Project)
                run2 = detail_para.add_run(value)
                run2.font.name = 'Calibri'
                run2.font.size = Pt(10)
                run2.font.bold = True
                run2.underline = True
        
        # Add spacing
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)  # Remove gap
        spacer.paragraph_format.space_after = Pt(0)
        
        # Dear Sir,
        dear_para = self.doc.add_paragraph('Dear Sir,')
        dear_para.paragraph_format.space_before = Pt(0)
        dear_para.paragraph_format.space_after = Pt(0)
        for run in dear_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
        
        # Introductory text
        intro_para = self.doc.add_paragraph()
        intro_para.paragraph_format.space_before = Pt(0)
        intro_para.paragraph_format.space_after = Pt(0)  # Remove gap
        run = intro_para.add_run('With reference to your enquiry, we would like to give our competitive offer for ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run = intro_para.add_run(self.subject)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run = intro_para.add_run(' as follows')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    
    def _create_header(self):
        """Create and format header row"""
        header_cells = self.table.rows[0].cells
        
        # Determine if discount is applied
        has_discount = getattr(self, 'has_discount', False)
        
        # Change header text based on discount status
        if has_discount:
            headers = ['SL.\nNO.', 'ITEM DESCRIPTION', 'UNIT', 'QTY', 'UNIT PRICE\n(AED)', 'DISCOUNTED TOTAL PRICE\n(AED)']
        else:
            headers = ['SL.\nNO.', 'ITEM DESCRIPTION', 'UNIT', 'QTY', 'UNIT PRICE\n(AED)', 'TOTAL PRICE\n(AED)']
        
        # Determine header color based on template
        if self.template_path.lower().endswith("colex_template.docx"):
            header_color = 'A3B463'  # Colex color
        else:
            header_color = '5F9EA0'  # Default color
        
        for idx, header_text in enumerate(headers):
            cell = header_cells[idx]
            cell.text = header_text
            # Apply header formatting
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), header_color)
            cell._element.get_or_add_tcPr().append(shading_elm)
            
            # Set vertical alignment to bottom
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'bottom')
            tcPr.append(vAlign)
            
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
    
    def _fill_common_row(self):
        """Fill the common row with preset phrases and common elements"""
        # Get company brand name from frontend or fallback to template-based detection
        if self.company_short_name:
            brand_name = self.company_short_name
        elif self.template_path.lower().endswith("colex_template.docx"):
            brand_name = "COLEX KOREA"
        else:
            brand_name = "PIPECO TANKS® MALAYSIA"
        
        # Build common text with brand name - reorganized into 3 lines
        # Line 1: Company name, warranty, and support system
        common_text = f"GRP SECTIONAL WATER TANK - 10 YEAR WARRANTY - {brand_name}"
        
        # Check if all tanks have the same support system and add to line 1
        common_support = self._get_common_support_system()
        if common_support:
            # All tanks have the same support system, add to common row
            support_text = self._get_support_system_text(common_support)
            common_text += f" - {support_text}"
        
        # Find common elements
        common_elements = self._find_common_elements()
        
        # Line 2: Tank type (insulation details)
        # Line 3: Skid base details
        for element_type, element_value in common_elements:
            if element_type == "type":
                # Add tank type on new line (line 2)
                common_text += f"\n{element_value}"
            elif element_type == "skid":
                # Add skid base on new line (line 3)
                common_text += f"\n{element_value}"
        
        # Set the text in the merged cell
        cell = self.table.rows[1].cells[0]
        cell.text = common_text
        
        # Make everything bold and Calibri 11, remove spacing
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
    
    def _fill_tank_row(self, row_idx, tank):
        """Fill a tank row with data"""
        row = self.table.rows[row_idx]
        
        # SL. NO. (center alignment)
        # Only set if this is the first option (option_number == 1)
        # For subsequent options, the cell will be merged later
        if tank.get('option_number', 1) == 1:
            cell = row.cells[0]
            cell.text = str(tank["sl_no"])
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set vertical alignment to center
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
        
        # Find common elements
        common_elements = self._find_common_elements()
        common_types = {elem[0] for elem in common_elements}
        
        # Item Description
        cell = row.cells[1]
        cell.text = ""  # Clear cell first
        paragraph = cell.paragraphs[0]
        
        # Set compact spacing for the paragraph
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0  # Single line spacing
        
        # Add tab stops for vertical alignment of colons and values
        # ADJUSTED: Colon position reduced to 1.05 inches, values at 1.25 inches
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(1.05))  # Position for colons
        tab_stops.add_tab_stop(Inches(1.25))  # Position for values after colon
        
        # Add OPTION label if there are multiple options for this tank
        option_total = tank.get('option_total', 1)
        option_roman = tank.get('option_roman', '')
        
        # Check if support systems are mixed across tanks
        common_support = self._get_common_support_system()
        
        if option_total > 1:
            # When there are options, check if we need to add support system on same line
            if common_support is None:
                # Mixed support systems - add support system text after OPTION on same line
                tank_support = tank.get('support_system', 'Internal')
                support_text = self._get_support_system_text(tank_support)
                run = paragraph.add_run(f"OPTION {option_roman} - {support_text}\n")
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
            else:
                # All same support system - just add OPTION label
                run = paragraph.add_run(f"OPTION {option_roman}\n")
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
        else:
            # No options - if mixed support, add support system text above tank name
            if common_support is None:
                tank_support = tank.get('support_system', 'Internal')
                support_text = self._get_support_system_text(tank_support)
                run = paragraph.add_run(f"{support_text}\n")
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
        
        # Tank name with underline and optionally skid in brackets
        tank_name = tank.get('name', '') or ''
        tank_skid = tank.get('skid', '') or ''
        
        if tank_name:
            # Check if skid is common
            skid_is_common = "skid" in common_types
            
            # Add partition status to tank name
            partition_status = " (WITH PARTITION)" if tank.get('partition', False) else ""
            
            if tank_skid and not skid_is_common:
                # Add tank name with partition status and underline
                run = paragraph.add_run(tank_name + partition_status)
                run.underline = True
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                
                # Add space and opening bracket with underline
                run = paragraph.add_run(" (")
                run.underline = True
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                
                # Add skid details with underline
                run = paragraph.add_run(tank_skid)
                run.underline = True
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                
                # Add closing bracket with underline, then no newline
                run = paragraph.add_run(")")
                run.underline = True
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
            else:
                run = paragraph.add_run(tank_name + partition_status)
                run.underline = True
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
        elif tank_skid and "skid" not in common_types:
            # No name, just print skid without brackets
            run = paragraph.add_run(tank_skid)
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
        
        # Type (only if not common) - with hanging indent for wrapped lines
        tank_type = tank.get('type', '') or ''
        if "type" not in common_types and tank_type:
            # Create separate paragraph for Type with hanging indent
            type_para = cell.add_paragraph()
            
            # Set compact spacing
            type_para.paragraph_format.space_before = Pt(0)
            type_para.paragraph_format.space_after = Pt(0)
            type_para.paragraph_format.line_spacing = 1.0  # Single line spacing
            
            # Set up tab stops for Type paragraph
            # ADJUSTED: Colon at 1.05 inches, value start at 1.25 inches
            type_tab_stops = type_para.paragraph_format.tab_stops
            type_tab_stops.add_tab_stop(Inches(1.05))  # Position for colon
            type_tab_stops.add_tab_stop(Inches(1.25))  # Position for value start
            
            # Configure hanging indent - continuation lines align with value start
            type_para.paragraph_format.left_indent = Inches(1.25)
            type_para.paragraph_format.first_line_indent = Inches(-1.25)
            
            run = type_para.add_run(f"Type\t:\t{tank_type}")
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
        
        # Add a new paragraph for Size and remaining items to keep them separate from Type
        size_para = cell.add_paragraph()
        
        # Set compact spacing
        size_para.paragraph_format.space_before = Pt(0)
        size_para.paragraph_format.space_after = Pt(0)
        size_para.paragraph_format.line_spacing = 1.0  # Single line spacing
        
        # Set up tab stops for size paragraph
        # ADJUSTED: Colon at 1.05 inches, values at 1.25 inches
        size_tab_stops = size_para.paragraph_format.tab_stops
        size_tab_stops.add_tab_stop(Inches(1.05))  # Position for colons
        size_tab_stops.add_tab_stop(Inches(1.25))  # Position for values after colon
        
        # Size (aligned format with tabs for colon alignment)
        # Handle potentially empty dimension values
        length_display = tank.get('length_display', '') or str(tank.get('length', '') or '')
        width_display = tank.get('width_display', '') or str(tank.get('width', '') or '')
        height_display = str(tank.get('height', '') or '')
        
        # Only show Size if at least one dimension exists
        if length_display or width_display or height_display:
            size_text = f"SIZE\t:\t{length_display} M (L) X {width_display} M (W) X {height_display} M (H)\n"
            run = size_para.add_run(size_text)
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
        
        # Total Capacity (aligned format with tabs for colon alignment)
        # Only show if volume exists (dimensions were provided)
        volume_m3 = tank.get('volume_m3', 0.0) or 0.0
        gallons = tank.get('gallons', 0.0) or 0.0
        
        if volume_m3 > 0:
            capacity_text = f"TOTAL CAPACITY\t:\t{volume_m3:.2f} M³ ({gallons:.0f} {self.gallon_type})"
            
            # Only show Free Board and Net Volume if needFreeBoard is enabled
            need_free_board = tank.get('need_free_board', False)
            if need_free_board:
                capacity_text += "\n"
            
            run = size_para.add_run(capacity_text)
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            
            # Show Net Volume and Free Board only if needFreeBoard is enabled
            if need_free_board:
                # Net Volume first (use pre-calculated value from tank data)
                net_volume_m3 = tank.get('net_volume_m3', volume_m3)
                # Calculate net volume gallons based on gallon type
                if self.gallon_type == "USG":
                    net_volume_gallons = net_volume_m3 * 264.172
                else:
                    net_volume_gallons = net_volume_m3 * 219.969
                run = size_para.add_run(f"NET VOLUME\t:\t{net_volume_m3:.2f} M³ ({net_volume_gallons:.0f} {self.gallon_type})\n")
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                
                # Free Board below Net Volume (aligned format with tabs for colon alignment)
                free_board_m = tank.get('free_board', 0.3)
                free_board_cm = free_board_m * 100  # Convert meters to cm
                run = size_para.add_run(f"FREE BOARD\t:\t{free_board_cm:.0f} cm ({free_board_m:.2f} M)")
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
        
        # Unit (center alignment both horizontal and vertical)
        cell = row.cells[2]
        cell.text = tank.get("unit", "") or ""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set vertical alignment to center
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
        
        # QTY (center alignment both horizontal and vertical)
        cell = row.cells[3]
        qty = tank.get("qty", 0) or 0
        if qty:
            cell.text = str(int(qty) if isinstance(qty, float) and qty.is_integer() else qty)
        else:
            cell.text = ""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set vertical alignment to center
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
        
        # Unit Price (right alignment horizontal, center vertical)
        cell = row.cells[4]
        unit_price = tank.get('unit_price', 0.0) or 0.0
        if unit_price:
            cell.text = f"{unit_price:,.2f}"
        else:
            cell.text = ""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Set vertical alignment to center
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
        
        # Total Price (right alignment horizontal, center vertical, make bold)
        cell = row.cells[5]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_price = tank.get('total_price', 0.0) or 0.0
        if total_price:
            run = paragraph.add_run(f"{total_price:,.2f}")
        else:
            run = paragraph.add_run("")
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        # Set vertical alignment to center
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
    
    def _add_ladder_row(self, row_idx):
        """Add ladder row if any tank is above 2M height"""
        row = self.table.rows[row_idx]
        
        # Find the highest sl_no
        max_sl_no = max(tank["sl_no"] for tank in self.tanks)
        
        row.cells[0].text = str(max_sl_no + 1)
        row.cells[1].text = "INTERNAL SS 316 AND EXTERNAL HDG SUPPORT SYSTEM "
        row.cells[2].text = "Set"
        row.cells[3].text = "1"
        row.cells[4].text = "0.00"
        row.cells[5].text = "0.00"
    
    def _create_footer(self, start_row):
        """Create footer rows with totals"""
        # Count how many total rows to show
        show_sub = getattr(self, 'show_sub_total', True)
        show_vat = getattr(self, 'show_vat', True)
        show_grand = getattr(self, 'show_grand_total', True)
        
        rows_to_show = sum([show_sub, show_vat, show_grand])
        
        if rows_to_show == 0:
            return  # Don't create footer if no totals to show
        
        # Merge SL. NO. and ITEM DESCRIPTION columns for footer rows
        self._merge_cells(start_row, 0, start_row + rows_to_show - 1, 1)
        
        # Remove left border and bottom border for the merged cell
        merged_cell = self.table.rows[start_row].cells[0]
        tc = merged_cell._element
        tcPr = tc.get_or_add_tcPr()
        
        # Remove existing borders
        for tcBorders in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(tcBorders)
        
        # Add cell borders with left and bottom set to 'none'
        tcBorders = OxmlElement('w:tcBorders')
        
        # Left border - none
        left_border = OxmlElement('w:left')
        left_border.set(qn('w:val'), 'none')
        left_border.set(qn('w:sz'), '0')
        left_border.set(qn('w:space'), '0')
        left_border.set(qn('w:color'), 'auto')
        tcBorders.append(left_border)
        
        # Bottom border - none (removes internal horizontal lines)
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'none')
        bottom_border.set(qn('w:sz'), '0')
        bottom_border.set(qn('w:space'), '0')
        bottom_border.set(qn('w:color'), 'auto')
        tcBorders.append(bottom_border)
        
        # Keep other borders normal (single)
        top_border = OxmlElement('w:top')
        top_border.set(qn('w:val'), 'single')
        top_border.set(qn('w:sz'), '4')
        top_border.set(qn('w:space'), '0')
        top_border.set(qn('w:color'), '000000')
        tcBorders.append(top_border)
        
        right_border = OxmlElement('w:right')
        right_border.set(qn('w:val'), 'single')
        right_border.set(qn('w:sz'), '4')
        right_border.set(qn('w:space'), '0')
        right_border.set(qn('w:color'), '000000')
        tcBorders.append(right_border)
        
        tcPr.append(tcBorders)
        
        # Calculate totals
        subtotal = sum(tank["total_price"] for tank in self.tanks)
        vat = subtotal * 0.05
        grand_total = subtotal + vat
        
        # Get flags for showing totals
        show_sub = getattr(self, 'show_sub_total', True)
        show_vat = getattr(self, 'show_vat', True)
        show_grand = getattr(self, 'show_grand_total', True)
        
        # Check if discount is applied
        has_discount = getattr(self, 'has_discount', False)
        
        current_row = start_row
        
        # Row 1: SUB TOTAL (if enabled)
        if show_sub:
            self._merge_cells(current_row, 2, current_row, 3)
            cell = self.table.rows[current_row].cells[2]
            # Change label based on discount status
            if has_discount:
                cell.text = 'DISCOUNTED SUB TOTAL:'
            else:
                cell.text = 'SUB TOTAL:'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            
            # Make AED bold and center aligned
            cell = self.table.rows[current_row].cells[4]
            cell.text = ""
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cell.paragraphs[0].add_run('AED')
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            
            # Make subtotal bold and right-aligned
            cell = self.table.rows[current_row].cells[5]
            cell.text = ""
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cell.paragraphs[0].add_run(f'{subtotal:,.2f}')
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            
            # Apply bold to label
            for paragraph in self.table.rows[current_row].cells[2].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
            
            current_row += 1
        
        # Row 2: VAT 5% (if enabled)
        if show_vat:
            self._merge_cells(current_row, 2, current_row, 3)
            cell = self.table.rows[current_row].cells[2]
            cell.text = 'VAT 5%:'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            
            # Make AED bold and center aligned
            cell = self.table.rows[current_row].cells[4]
            cell.text = ""
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cell.paragraphs[0].add_run('AED')
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            
            # Make VAT bold and right-aligned
            cell = self.table.rows[current_row].cells[5]
            cell.text = ""
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cell.paragraphs[0].add_run(f'{vat:,.2f}')
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            
            # Apply bold to label
            for paragraph in self.table.rows[current_row].cells[2].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
            
            current_row += 1
        
        # Row 3: GRAND TOTAL (if enabled)
        if show_grand:
            self._merge_cells(current_row, 2, current_row, 3)
            cell = self.table.rows[current_row].cells[2]
            cell.text = 'GRAND TOTAL:'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            
            # Make AED bold and center aligned
            cell = self.table.rows[current_row].cells[4]
            cell.text = ""
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cell.paragraphs[0].add_run('AED')
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            
            # Make grand total bold and right-aligned
            cell = self.table.rows[current_row].cells[5]
            cell.text = ""
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cell.paragraphs[0].add_run(f'{round(grand_total):,.0f}')
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            
            # Apply bold to label
            for paragraph in self.table.rows[current_row].cells[2].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
    
    def _merge_cells(self, row, start_col, end_row, end_col):
        """Merge cells in the table"""
        cell = self.table.cell(row, start_col)
        end_cell = self.table.cell(end_row, end_col)
        cell.merge(end_cell)
    
    def _apply_font_to_all_cells(self):
        """Apply Calibri 11 font to all cells"""
        for row in self.table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)
    
    def _apply_table_borders(self):
        """Manually apply borders to the table when Table Grid style is not available"""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # Define border style
        borders_xml = (
            '<w:tblBorders %s>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tblBorders>'
        ) % nsdecls('w')
        
        # Apply borders to table
        tbl = self.table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml('<w:tblPr %s/>' % nsdecls('w'))
            tbl.insert(0, tblPr)
        
        # Remove existing borders if any
        for borders in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(borders)
        
        # Add new borders
        tblPr.append(parse_xml(borders_xml))
    
    def _remove_cell_padding(self):
        """Remove cell padding/margins to make text compact"""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # Set cell margins to 0 (or very small values)
        cell_margins_xml = (
            '<w:tblCellMar %s>'
            '<w:top w:w="0" w:type="dxa"/>'
            '<w:left w:w="30" w:type="dxa"/>'  # Small left margin for readability
            '<w:bottom w:w="0" w:type="dxa"/>'
            '<w:right w:w="30" w:type="dxa"/>'  # Small right margin for readability
            '</w:tblCellMar>'
        ) % nsdecls('w')
        
        # Apply to table properties
        tbl = self.table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml('<w:tblPr %s/>' % nsdecls('w'))
            tbl.insert(0, tblPr)
        
        # Remove existing cell margins if any
        for cellMar in tblPr.findall(qn('w:tblCellMar')):
            tblPr.remove(cellMar)
        
        # Add new cell margins
        tblPr.append(parse_xml(cell_margins_xml))
    
    def _apply_column_specific_padding(self):
        """Apply specific padding to certain columns: 1pt space before description, 1pt after price columns"""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # 1 point = 20 dxa (twentieths of a point)
        for row in self.table.rows:
            # Column 1 (ITEM DESCRIPTION) - add 1pt left margin
            cell = row.cells[1]
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            
            # Remove existing cell margins
            for cellMar in tcPr.findall(qn('w:tcMar')):
                tcPr.remove(cellMar)
            
            # Add 1pt (20 dxa) left margin for description
            cell_mar_xml = (
                '<w:tcMar %s>'
                '<w:left w:w="20" w:type="dxa"/>'
                '</w:tcMar>'
            ) % nsdecls('w')
            tcPr.append(parse_xml(cell_mar_xml))
            
            # Column 4 (UNIT PRICE) - add 3pt right margin
            cell = row.cells[4]
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            
            # Remove existing cell margins
            for cellMar in tcPr.findall(qn('w:tcMar')):
                tcPr.remove(cellMar)
            
            # Add 3pt (60 dxa) right margin
            cell_mar_xml = (
                '<w:tcMar %s>'
                '<w:right w:w="60" w:type="dxa"/>'
                '</w:tcMar>'
            ) % nsdecls('w')
            tcPr.append(parse_xml(cell_mar_xml))
            
            # Column 5 (TOTAL PRICE) - add 3pt right margin
            cell = row.cells[5]
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            
            # Remove existing cell margins
            for cellMar in tcPr.findall(qn('w:tcMar')):
                tcPr.remove(cellMar)
            
            # Add 3pt (60 dxa) right margin
            cell_mar_xml = (
                '<w:tcMar %s>'
                '<w:right w:w="60" w:type="dxa"/>'
                '</w:tcMar>'
            ) % nsdecls('w')
            tcPr.append(parse_xml(cell_mar_xml))
    
    def _add_additional_sections(self):
        """Add all additional sections after the quotation table"""
        if self.sections.get('note'):
            self._add_note_section()
        
        if self.sections.get('closing'):
            self._add_closing_paragraph()
        
        if self.sections.get('signature'):
            self._add_signature_section()
        
        if self.sections.get('material_spec'):
            self._add_material_spec_section()
        
        if self.sections.get('warranty'):
            self._add_warranty_section()
        
        if self.sections.get('terms'):
            self._add_terms_section()
        
        if self.sections.get('supplier_scope'):
            self._add_supplier_scope_section()
        
        if self.sections.get('customer_scope'):
            self._add_customer_scope_section()
        
        if self.sections.get('extra_note'):
            self._add_extra_note_section()
        
        if self.sections.get('final_note'):
            self._add_final_note_section()
        
        if self.sections.get('thank_you'):
            self._add_thank_you_section()
    
    def _add_note_section(self):
        """Add NOTE section as numbered list with heading 'Note:'"""
        # Add adjustable spacing before NOTE section
        if self.note_section_gap > 0:
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(self.note_section_gap)
            spacer.paragraph_format.space_after = Pt(0)

        # Add heading 'NOTE:'
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run('NOTE:')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True

        # Add each note as a numbered point - NOW WITH BOLD TEXT
        notes = self.section_content['note']
        for idx, note in enumerate(notes, 1):
            para = self.doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(f"{idx}. {note.strip()}")
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.font.bold = True  # Changed from False to True
    
    def _add_closing_paragraph(self):
        """Add closing paragraph"""
        # Add adjustable spacing before closing paragraph
        if self.closing_paragraph_gap > 0:
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(self.closing_paragraph_gap)
            spacer.paragraph_format.space_after = Pt(0)
        
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(self.section_content['closing'])
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    
    def _add_signature_section(self):
        """Add signature section with Yours truly and signatories"""
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(12)
        spacer.paragraph_format.space_after = Pt(0)
        
        # Yours truly - keep with next to prevent page break
        para = self.doc.add_paragraph('Yours truly,')
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.keep_with_next = True
        para.paragraph_format.keep_together = True
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
        
        # For GRP PIPECO TANKS TRADING L.L.C - keep with next
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.keep_with_next = True
        para.paragraph_format.keep_together = True
        run = para.add_run(f'For {self._get_company_name()}')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.italic = True
        
        # Add signature image if provided - keep with next
        if self.section_content['signature'].get('signature_image'):
            try:
                para = self.doc.add_paragraph()
                para.paragraph_format.keep_with_next = True
                para.paragraph_format.keep_together = True
                run = para.add_run()
                # Signature image size reduced from 2 inches to 1.2 inches for better appearance
                run.add_picture(self.section_content['signature']['signature_image'], width=Inches(0.9))
            except:
                # Add placeholder text if image not found
                para = self.doc.add_paragraph('[Signature Image]')
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.keep_with_next = True
                para.paragraph_format.keep_together = True
        else:
            # Add space for signature - keep with next
            for _ in range(3):
                para = self.doc.add_paragraph()
                para.paragraph_format.keep_with_next = True
                para.paragraph_format.keep_together = True
        
        # Create 2-column layout table for signatories (matching the first table width)
        sig_table = self.doc.add_table(rows=1, cols=2)
        sig_table.autofit = False
        sig_table.columns[0].width = Inches(5.0)  # Match left column of client info table
        sig_table.columns[1].width = Inches(2.5)  # Match right column of client info table
        
        # Add left indent to match the client info table position
        tbl = sig_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Set table indent
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '72')  # Reduced indent: 72 twips = 0.05 inches (smaller left indent)
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        
        # Remove table borders
        for borders in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(borders)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Remove cell borders and padding for each cell
        row = sig_table.rows[0]
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            
            # Remove existing borders
            for tcBorders in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(tcBorders)
            
            # Add cell borders set to 'none'
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)
            tcPr.append(tcBorders)
            
            # Remove cell padding
            for tcMar in tcPr.findall(qn('w:tcMar')):
                tcPr.remove(tcMar)
            tcMar = OxmlElement('w:tcMar')
            for margin_name in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{margin_name}')
                margin.set(qn('w:w'), '0')
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)
            tcPr.append(tcMar)
        
        # Left signatory
        left_cell = row.cells[0]
        # Set cell width using XML (6900 twips = 5.0 inches)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), '7200')
        tcW.set(qn('w:type'), 'dxa')
        left_cell._element.get_or_add_tcPr().append(tcW)
        
        self._add_signatory_info_table(left_cell, 
                                self.section_content['signature']['left_name'],
                                self.section_content['signature']['left_title'],
                                self.section_content['signature']['left_mobile'],
                                self.section_content['signature']['left_email'])
        
        # Right signatory
        right_cell = row.cells[1]
        # Set cell width using XML (3900 twips = 2.5 inches) 
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), '3600')
        tcW.set(qn('w:type'), 'dxa')
        tcPr = right_cell._element.get_or_add_tcPr()
        tcPr.append(tcW)
        
        # Set vertical alignment to top
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'top')
        tcPr.append(vAlign)
        
        self._add_signatory_info_table(right_cell,
                                self.section_content['signature']['right_name'],
                                self.section_content['signature']['right_title'],
                                self.section_content['signature']['right_mobile'],
                                self.section_content['signature']['right_email'])
        
        # Add seal if provided (outside the table, floating to the right)
        # Seal image insertion removed as requested
    
    def _add_signatory_info_table(self, cell, name, title, mobile, email):
        """Add signatory information to a cell in table format"""
        # Remove default paragraph
        if len(cell.paragraphs) > 0:
            p = cell.paragraphs[0]._element
            p.getparent().remove(p)
        
        # Track if any content was added
        content_added = False
        
        # Name - only if not empty
        if name:
            para = cell.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(name)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
            content_added = True
        
        # Title - only if not empty
        if title:
            para = cell.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(title)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
            content_added = True
        
        # Mobile - only if not empty
        if mobile:
            para = cell.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(f'MOB: {mobile}')
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
            content_added = True
        
        # Email - only if not empty
        if email:
            para = cell.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run('EMAIL: ')
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
            run2 = para.add_run(email)
            run2.font.name = 'Calibri'
            run2.font.size = Pt(10)
            run2.font.color.rgb = RGBColor(0, 0, 255)
            run2.font.underline = True
            content_added = True
        
        # If no content was added, ensure cell has at least one empty paragraph
        # (Word requires every cell to have at least one paragraph element)
        if not content_added:
            para = cell.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
    
    def _add_material_spec_section(self):
        """Add material specification section"""
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(12)
        spacer.paragraph_format.space_after = Pt(0)
        
        # Heading
        heading = self.doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)
        run = heading.add_run('MATERIAL SPECIFICATION: -')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Add each spec
        for spec in self.section_content['material_spec']:
            para = self.doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            run = para.add_run(f'➢  {spec}')
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
    
    def _add_warranty_section(self):
        """Add warranty section"""
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(12)
        spacer.paragraph_format.space_after = Pt(0)
        
        # Heading
        heading = self.doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)
        run = heading.add_run('THE WARRANTY WILL NOT BE APPLICABLE FOR THE FOLLOWING CASES:')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Add each item
        for item in self.section_content['warranty']:
            para = self.doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            run = para.add_run(f'➢  {item}')
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
    
    def _add_terms_section(self):
        """Add terms and conditions section"""
        # DEBUG
        print(f"\n>>> INSIDE _add_terms_section()")
        print(f">>>   section_content['terms'] exists: {'terms' in self.section_content}")
        if 'terms' in self.section_content:
            print(f">>>   Number of formatted terms: {len(self.section_content['terms'])}")
            print(f">>>   Terms keys: {list(self.section_content['terms'].keys())}")
        if 'terms_plain' in self.section_content:
            print(f">>>   Number of plain terms: {len(self.section_content['terms_plain'])}")
        
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)
        
        # Heading
        heading = self.doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)
        run = heading.add_run('TERMS AND CONDITIONS: -')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Add each formatted term (key: value) with specific spacing before colon for alignment
        spacing_map = {
            'Price': '         ',      # 9 spaces
            'Validity': '    ',         # 4 spaces
            'Delivery': '    ',         # 4 spaces
            'Payment': '  '             # 2 spaces
        }
        
        terms_added = 0
        for key, value in self.section_content['terms'].items():
            para = self.doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            
            # Get spacing for this key, default to 2 spaces if not in map
            spacing = spacing_map.get(key, '  ')
            
            # Add key with colon and value on same line
            run = para.add_run(f'➢  {key}{spacing}: {value}')
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            terms_added += 1
        
        # Add plain custom terms (without key: value format)
        if 'terms_plain' in self.section_content:
            for plain_term in self.section_content['terms_plain']:
                if plain_term and plain_term.strip():  # Only add non-empty terms
                    para = self.doc.add_paragraph()
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.first_line_indent = Inches(-0.25)
                    
                    run = para.add_run(f'➢  {plain_term}')
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
                    terms_added += 1
        
        print(f">>>   Terms actually added to document: {terms_added}")
        print(f">>> EXITING _add_terms_section()\n")
    
    def _add_extra_note_section(self):
        """Add extra NOTE section"""
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(12)
        spacer.paragraph_format.space_after = Pt(0)
        
        # Heading
        heading = self.doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)
        run = heading.add_run('NOTE:')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Add each item with bullet points, bold company name if present
        company_name = self._get_company_name()
        for item in self.section_content['extra_note']:
            para = self.doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            if company_name in item:
                before, after = item.split(company_name, 1)
                run = para.add_run('➢  ' + before)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run = para.add_run(company_name)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.bold = True
                run = para.add_run(after)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
            else:
                run = para.add_run(f'➢  {item}')
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
    
    def _add_supplier_scope_section(self):
        """Add supplier scope section"""
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(12)
        spacer.paragraph_format.space_after = Pt(0)
        
        # Heading
        heading = self.doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)
        run = heading.add_run('SUPPLIER SCOPE: -')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Add each item
        for item in self.section_content['supplier_scope']:
            para = self.doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            run = para.add_run(f'➢  {item}')
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
    
    def _add_customer_scope_section(self):
        """Add customer scope section"""
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(12)
        spacer.paragraph_format.space_after = Pt(0)
        
        # Heading
        heading = self.doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)
        run = heading.add_run('CUSTOMER SCOPE: -')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Add each item
        for item in self.section_content['customer_scope']:
            para = self.doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            run = para.add_run(f'➢  {item}')
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
    
    def _add_final_note_section(self):
        """Add final note section"""
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(12)
        spacer.paragraph_format.space_after = Pt(0)
        
        # Heading
        heading = self.doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)
        run = heading.add_run('NOTE: -')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Add each note
        company_name = self._get_company_name()
        for note in self.section_content['final_note']:
            para = self.doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            
            # Check if this note contains the company name
            if company_name in note:
                # Split the note by company name and add with bold highlighting
                parts = note.split(company_name)
                run = para.add_run(f'➢  {parts[0]}')
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                
                # Add company name in bold
                run = para.add_run(company_name)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.bold = True
                
                # Add remaining text
                run = para.add_run(parts[1])
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
            else:
                run = para.add_run(f'➢  {note}')
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
    
    def _add_thank_you_section(self):
        """Add thank you section"""
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(24)
        spacer.paragraph_format.space_after = Pt(0)
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run('THANK YOU FOR YOUR BUSINESS')
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 32, 96)
    
    def add_additional_pages(self, num_pages):
        """Add additional pages with header table (for pages 2 onwards)"""
        for page_num in range(2, num_pages + 1):
            # Add page break
            self.doc.add_page_break()
            
            # Add header table for this page
            print(f"  → Adding header table to page {page_num}")
            self._add_header_table_for_additional_pages()
            
            # Add a placeholder paragraph to ensure content exists on the page
            placeholder = self.doc.add_paragraph()
            placeholder.paragraph_format.space_before = Pt(12)
            placeholder.paragraph_format.space_after = Pt(0)
            placeholder.add_run(f"[Page {page_num}]")  # Visible marker for debugging
    
    def remove_document_protection(self):
        """Remove all document protection and read-only settings to make document fully editable"""
        try:
            # Access the document settings part
            settings_element = self.doc.settings.element
            
            # Remove write protection (read-only recommendation)
            write_protection = settings_element.find(qn('w:writeProtection'))
            if write_protection is not None:
                settings_element.remove(write_protection)
                print("✓ Removed write protection from document")
            
            # Remove document protection
            doc_protection = settings_element.find(qn('w:documentProtection'))
            if doc_protection is not None:
                settings_element.remove(doc_protection)
                print("✓ Removed document protection")
                
        except Exception as e:
            print(f"⚠ Warning: Could not remove document protection: {e}")
            # Continue anyway - document will still be saved
    
    def save(self, filename=None):
        """Save the document - will replace existing file with same name"""
        if filename is None:
            # Use quote number for filename, replacing invalid characters
            quote_no = getattr(self, 'quote_number', 'quote')
            # Replace invalid filename characters with underscores
            safe_quote_no = quote_no.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
            filename = f"{safe_quote_no}.docx"
        
        # Check if filename is already an absolute path
        if os.path.isabs(filename):
            # Use the provided absolute path directly
            full_path = filename
            # Create parent directory if it doesn't exist
            os.makedirs(os.path.dirname(full_path), exist_ok=True)
        else:
            # Create Final_Doc folder if it doesn't exist
            final_doc_folder = os.path.join(os.path.dirname(__file__), "Final_Doc")
            os.makedirs(final_doc_folder, exist_ok=True)
            
            # Build full path to save in Final_Doc folder
            full_path = os.path.join(final_doc_folder, filename)
        
        # Check if file exists and notify user about replacement
        if os.path.exists(full_path):
            print(f"\n⚠ File '{os.path.basename(full_path)}' already exists - replacing with new version...")
        
        # Remove any document-level protection and read-only settings before saving
        self.remove_document_protection()
        
        # Save document (will overwrite existing file)
        self.doc.save(full_path)
        
        # Remove read-only attribute to ensure file is editable on server
        try:
            # Set file permissions to be readable and writable (remove read-only)
            current_permissions = os.stat(full_path).st_mode
            os.chmod(full_path, current_permissions | stat.S_IWRITE | stat.S_IREAD)
            print(f"✓ File permissions set to editable mode")
        except Exception as perm_error:
            print(f"⚠ Warning: Could not set file permissions: {perm_error}")
            # Continue anyway - file is saved, just might have permission issues
        
        print(f"✓ Document saved as: {full_path}")
        print(f"✓ Template content preserved with table added")
        return full_path

def main():
    """Main function to run the generator"""
    try:
        # Template selection
        print("\n" + "="*60)
        print("TEMPLATE SELECTION")
        print("="*60)
        print("1. GRP Tanks")
        print("2. Pipeco")
        print("3. Colex")
        
        template_choice = input("Select template (1/2/3): ").strip()
        template_map = {
            "1": "grp_template.docx",
            "2": "pipeco_template.docx",
            "3": "colex_template.docx"
        }
        template_path = template_map.get(template_choice, "grp_template.docx")
        
        generator = TankInvoiceGenerator(template_path=template_path)
        generator.get_user_inputs()
        generator.create_invoice_table()
        
        # Delete quote box from first page header only
        generator.delete_tables_in_first_page_header()
        
        generator.save()
        
        print("\n" + "="*60)
        print("Invoice generated successfully!")
        print("="*60)
        
    except KeyboardInterrupt:
        print("\n\nProcess cancelled by user.")
    except Exception as e:
        print(f"\nError occurred: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
